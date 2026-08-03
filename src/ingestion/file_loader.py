from pathlib import Path
import logging

from src.ingestion.document import Document, make_doc_id

logger = logging.getLogger(__name__)

class FileLoader:
    def __init__(self, directory):
        self.directory = Path(directory)
    
    def load(self):
        for file in self.directory.iterdir():
            try:
                if file.suffix == '.txt':
                    text = file.read_text(encoding = 'utf-8')
                elif file.suffix == '.pdf':
                    text = self._extract_pdf(file)
                elif file.suffix == '.docx':
                    text = self._extract_docx(file)
                else:
                    continue
                
                if not text.strip():
                    logger.warning(f'file {file.name} in file {file} is extracted but its empty')
                    continue
                
                yield Document(
                    doc_id = make_doc_id(file.name + text[:100]),
                    source = f'file : {file.name}',
                    text = text,
                    metadata = {'filetype':file.suffix.lstrip(".")}
                    
                )
            except Exception as e:
                logger.error(f'Failed to load {file.name}: due to error: {e}')
                
    def _extract_pdf(self, file):
        import pypdf
        text_parts = []
        with open(file,'rb') as f:
            reader= pypdf.PdfReader(f)
            for page in reader.pages:
                text_parts.append(page.extract_text())
            return '\n'.join(text_parts)
    
    def _extract_docx(self,file):
        import docx
        doc = docx.Document(file)
        return '\n'.join(p.text for p in doc.paragraphs)
        
            
            
            